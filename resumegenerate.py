from linkedin_scraper import Person, actions
from selenium import webdriver
from selenium.webdriver.chrome.service import Service as ChromeService
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
import time

# Function to scrape LinkedIn profile and generate basic resume
def generate_resume(linkedin_url, username, password):
    # Set up Selenium WebDriver (Chrome)
    service = ChromeService(ChromeDriverManager().install())
    driver = webdriver.Chrome(service=service)
    
    try:
        print("Logging into LinkedIn...")
        actions.login(driver, username, password)  # LinkedIn login
        time.sleep(5)  # Wait for login to complete

        print("Accessing LinkedIn profile...")
        # Scrape LinkedIn profile
        person = Person(linkedin_url, driver=driver)
        person.scrape(close_on_complete=False)
        
        print("Creating resume document...")
        # Create a new Word document for the resume
        doc = Document()
        doc.add_heading('Resume', 0)

        # Add basic sections
        doc.add_heading('Name', level=1)
        doc.add_paragraph(person.name)
        
        doc.add_heading('Title', level=1)
        doc.add_paragraph(person.job_title)
        
        doc.add_heading('Contact Information', level=1)
        doc.add_paragraph(f'LinkedIn: {linkedin_url}')
        
        doc.add_heading('Experience', level=1)
        for exp in person.experiences:
            doc.add_heading(exp.job_title, level=2)
            doc.add_paragraph(f'Company: {exp.company_name}')
            doc.add_paragraph(f'Dates: {exp.date_range}')
            doc.add_paragraph(f'Description: {exp.description}')

        doc.add_heading('Education', level=1)
        for edu in person.education:
            doc.add_heading(edu.institution_name, level=2)
            doc.add_paragraph(f'Degree: {edu.degree}')
            doc.add_paragraph(f'Dates: {edu.date_range}')
            doc.add_paragraph(f'Description: {edu.description}')

        doc.add_heading('Skills', level=1)
        doc.add_paragraph(", ".join(person.skills))

        doc_path = 'basic_resume.docx'
        doc.save(doc_path)
        
        return doc_path
    except Exception as e:
        print(f"An error occurred: {e}")
    finally:
        driver.quit()

def add_additional_section(doc_path, section_title, section_content):
    doc = Document(doc_path)
    doc.add_heading(section_title, level=1)
    doc.add_paragraph(section_content)
    doc.save(doc_path)

if __name__ == "__main__":
    linkedin_url = input("Enter your LinkedIn profile URL: ")
    username = input("Enter your LinkedIn username: ")
    password = input("Enter your LinkedIn password: ")

    doc_path = generate_resume(linkedin_url, username, password)
    if doc_path:
        print(f"Basic resume generated: {doc_path}")

        while True:
            add_section = input("Do you want to add any additional sections? (yes/no): ")
            if add_section.lower() == 'yes':
                section_title = input("Enter the title of the additional section: ")
                section_content = input("Enter the content of the additional section: ")
                add_additional_section(doc_path, section_title, section_content)
                print(f"Section '{section_title}' added to the resume.")
            else:
                break

        print("Resume update completed.")